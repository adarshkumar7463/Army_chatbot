import os
import uuid
from datetime import datetime

from django.conf import settings
import csv
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

# Ensure export folder exists
EXPORT_DIR = os.path.join(settings.MEDIA_ROOT, "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


def export_officers(officers, title, export_type):
    headers = ["Army Number", "Full Name", "Rank", "Unit", "DOB", "Phone", "Email"]
    rows = [
        [
            o.army_number,
            o.full_name,
            o.rank,
            o.unit,
            o.dob.strftime("%d-%m-%Y") if o.dob else "",
            o.phone,
            o.email,
        ]
        for o in officers
    ]
    return _generate_file(headers, rows, title, export_type, "officers")


def export_family(families, title, export_type):
    headers = ["Officer Army No", "Relation", "Name", "Age"]
    rows = [
        [f.officer.army_number, f.relation, f.name, f.contact] for f in families
    ]
    return _generate_file(headers, rows, title, export_type, "family")


def export_education(educations, title, export_type):
    headers = ["Officer Army No", "Degree", "Institution", "Year"]
    rows = [
        [e.officer.army_number, e.degree, e.institution, e.year_of_passing]
        for e in educations
    ]
    return _generate_file(headers, rows, title, export_type, "education")


def export_awards(awards, title, export_type):
    headers = ["Officer Army No", "Award", "Reason", "Location", "Date"]
    rows = [
        [
            a.officer.army_number,
            a.award_name,
            a.reason,
            a.location,
            a.date_awarded.strftime("%d-%m-%Y") if a.date_awarded else "",
        ]
        for a in awards
    ]
    return _generate_file(headers, rows, title, export_type, "awards")


# ---------- Core generator ----------
def _generate_file(headers, rows, title, export_type, filename_base):
    # Unique filename with timestamp + random ID
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = uuid.uuid4().hex[:6]
    filename = f"{filename_base}_{timestamp}_{unique_id}"

    if export_type == "excel":  # CSV export
        file_path = os.path.join(EXPORT_DIR, f"{filename}.csv")
        with open(file_path, mode="w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([title])
            writer.writerow(headers)
            writer.writerows(rows)

    elif export_type == "word":
        file_path = os.path.join(EXPORT_DIR, f"{filename}.docx")
        doc = Document()
        doc.add_heading("Army Record System", 0)
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        doc.save(file_path)

    elif export_type == "pdf":
        file_path = os.path.join(EXPORT_DIR, f"{filename}.pdf")
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(file_path)
        elements = [
            Paragraph("Army Record System", styles["Heading1"]),
            Paragraph(title, styles["Heading2"]),
            Spacer(1, 12),
        ]
        data = [headers] + rows
        table = Table(data)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ]
            )
        )
        elements.append(table)
        doc.build(elements)

    else:
        raise ValueError("Unsupported export type")

    # ✅ Return public URL for chatbot
    return settings.MEDIA_URL + "exports/" + os.path.basename(file_path)



def export_single_officer(officer, export_type, unique_id=""):
    """Export full details of a single officer (basic + family + education + awards)"""
    title = f"Officer {officer.army_number} - Full Details"
    headers = ["Section", "Field", "Value"]
    rows = []

    # Officer basic
    rows += [
        ["Officer", "Army Number", officer.army_number],
        ["Officer", "Full Name", officer.full_name],
        ["Officer", "Rank", officer.rank],
        ["Officer", "Unit", officer.unit],
        ["Officer", "DOB", officer.dob.strftime("%d-%m-%Y") if officer.dob else ""],
        ["Officer", "Phone", officer.phone],
        ["Officer", "Email", officer.email],
    ]

    # Family
    families = Family.objects.filter(officer=officer)
    for f in families:
        rows.append(["Family", f.relation, f.name])

    # Education
    educations = Education.objects.filter(officer=officer)
    for e in educations:
        rows.append(["Education", e.degree, f"{e.institution} ({e.year_of_passing})"])

    # Awards
    awards = Award.objects.filter(officer=officer)
    for a in awards:
        rows.append([
            "Award", a.award_name, 
            f"{a.reason} at {a.location} ({a.date_awarded.strftime('%d-%m-%Y') if a.date_awarded else ''})"
        ])

    return _generate_file(headers, rows, title, export_type, f"officer_{officer.army_number}_{unique_id}")
